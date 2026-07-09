from __future__ import annotations

import base64
import hashlib
import json
import math
import mimetypes
import os
import re
import sqlite3
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pymysql
import requests
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RAG_ROOT = PROJECT_ROOT / "rag"
load_dotenv(RAG_ROOT / ".env")

TEXT_SUFFIXES = {
    ".css",
    ".html",
    ".java",
    ".json",
    ".jsp",
    ".md",
    ".ps1",
    ".py",
    ".sql",
    ".ts",
    ".tsx",
    ".txt",
    ".xml",
}
SKIP_DIRS = {
    ".git",
    ".idea",
    ".venv",
    "__pycache__",
    "assets",
    "dist",
    "node_modules",
    "target",
    "uploads",
}
MAX_SOURCE_FILE_BYTES = 220_000
DEFAULT_CHAT_MODEL = "glm-4-flash"
DEFAULT_EMBED_MODEL = "embedding-3"
DEFAULT_VISION_MODEL = "glm-4v-flash"
LAST_EMBEDDING_ERROR = ""
ALLOWED_ORIGINS = [
    item.strip()
    for item in os.getenv("RAG_ALLOWED_ORIGINS", "*").split(",")
    if item.strip()
] or ["*"]
TOKEN_ALIASES = {
    "后台": ["admin", "manage", "management"],
    "管理": ["admin", "manager"],
    "管理端": ["admin"],
    "前台": ["client", "frontend", "home"],
    "职位": ["job", "jobs"],
    "岗位": ["job", "jobs"],
    "企业": ["company", "companies"],
    "用户": ["user", "users"],
    "简历": ["resume", "resumes"],
    "附件": ["document", "documents"],
    "上传": ["upload"],
    "投递": ["application", "applications", "apply"],
    "申请": ["application", "applications", "apply"],
    "地图": ["map", "amap"],
    "热力图": ["heatmap", "map"],
    "问答": ["rag", "chat"],
    "识图": ["vision", "image"],
}


class ChatRequest(BaseModel):
    question: str = Field(min_length=1, max_length=1000)
    topK: int = Field(default=6, ge=1, le=12)
    userContext: dict[str, Any] | None = None


class Source(BaseModel):
    id: str
    title: str
    type: str = "document"
    source: str
    snippet: str
    score: float = 0


class ChatResponse(BaseModel):
    answer: str
    sources: list[Source]
    mode: str
    configured: bool


class JobMatch(BaseModel):
    jobId: int | None = None
    title: str
    company: str = ""
    city: str = ""
    category: str = ""
    salary: str = ""
    score: float = 0
    reason: str = ""


class ResumeAnalysisResponse(BaseModel):
    fileName: str
    extractedText: str
    summary: str
    profile: dict[str, Any]
    skills: list[str]
    directions: list[str]
    matches: list[JobMatch]
    suggestions: list[str]
    mode: str
    configured: bool


class VisionAnalyzeResponse(BaseModel):
    fileName: str
    answer: str
    mode: str
    configured: bool
    sources: list[Source] = Field(default_factory=list)


@dataclass
class Document:
    id: str
    title: str
    text: str
    type: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)
    terms: set[str] = field(default_factory=set)
    text_hash: str = ""
    embedding: list[float] | None = None


app = FastAPI(title="Q_ITOffer RAG Service", version="1.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

INDEX: list[Document] = []


@app.on_event("startup")
def startup() -> None:
    rebuild_index(build_embeddings=False)


@app.get("/health")
def health() -> dict[str, Any]:
    embedded = sum(1 for doc in INDEX if doc.embedding)
    configured = bool(api_key())
    return {
        "status": "ok",
        "configured": configured,
        "chatConfigured": configured and bool(chat_model()),
        "embeddingConfigured": configured and bool(embed_model()),
        "visionConfigured": configured and bool(vision_model()),
        "documents": len(INDEX),
        "embeddedDocuments": embedded,
        "database": db_configured(),
        "embeddingError": LAST_EMBEDDING_ERROR,
        "retrievalVersion": 2,
    }


@app.get("/sources")
def sources() -> dict[str, Any]:
    items = [
        {
            "id": doc.id,
            "title": doc.title,
            "type": doc.type,
            "source": doc.source,
            "snippet": compact(doc.text, 150),
        }
        for doc in INDEX[:300]
    ]
    return {"items": items, "total": len(INDEX)}


@app.post("/index/rebuild")
def index_rebuild() -> dict[str, Any]:
    rebuild_index(build_embeddings=True)
    return {
        "status": "ok",
        "documents": len(INDEX),
        "embeddedDocuments": sum(1 for doc in INDEX if doc.embedding),
    }


@app.post("/chat", response_model=ChatResponse)
def chat(payload: ChatRequest) -> ChatResponse:
    user_docs = user_context_documents(payload.userContext)
    if needs_private_context(payload.question) and not user_docs:
        hits = retrieve(payload.question, min(payload.topK, 4))
        answer = (
            "当前还没有登录，或者没有拿到当前登录人的简历/投递上下文。"
            "我不能使用其他用户的简历和投递记录来判断你的情况。"
            "你可以先登录并完善简历，或者直接上传简历文件做匹配；在未登录状态下，我只能基于公共岗位和项目资料回答。"
        )
        return ChatResponse(answer=answer, sources=[source_from_hit(item) for item in hits], mode="fallback", configured=bool(api_key()))

    hits = retrieve(payload.question, payload.topK, user_docs)
    if is_personal_job_question(payload.question) and user_docs:
        answer = personal_job_answer(user_docs)
        return ChatResponse(answer=answer, sources=[source_from_hit(item) for item in hits], mode="fallback", configured=bool(api_key()))
    if not hits:
        return ChatResponse(
            answer="当前索引里没有找到足够相关的信息。可以先刷新索引，或确认数据库、README、docs 和源码目录是否已经准备好。",
            sources=[],
            mode="fallback",
            configured=bool(api_key()),
        )

    if api_key():
        try:
            answer = call_llm(payload.question, hits, bool(user_docs))
            return ChatResponse(answer=answer, sources=[source_from_hit(item) for item in hits], mode="llm", configured=True)
        except requests.RequestException as exc:
            fallback = fallback_answer(payload.question, hits)
            fallback += f"\n\n大模型接口暂时不可用，已切换本地检索回答。错误：{exc}"
            return ChatResponse(answer=fallback, sources=[source_from_hit(item) for item in hits], mode="fallback", configured=True)

    return ChatResponse(
        answer=fallback_answer(payload.question, hits),
        sources=[source_from_hit(item) for item in hits],
        mode="fallback",
        configured=False,
    )


@app.post("/resume/analyze", response_model=ResumeAnalysisResponse)
async def analyze_resume(file: UploadFile = File(...), topK: int = Form(8)) -> ResumeAnalysisResponse:
    data = await checked_upload(file)
    extracted_text, extract_mode = extract_resume_text(file.filename or "resume", file.content_type or "", data)
    if not extracted_text.strip() and is_image_file(file.filename or "", file.content_type or ""):
        extracted_text = vision_extract_resume(file.filename or "resume-image", file.content_type or "image/png", data)
        extract_mode = "vision" if extracted_text else "fallback"

    profile = extract_resume_profile(extracted_text)
    skills = profile.get("skills", [])
    directions = infer_directions(extracted_text, skills)
    matches = match_resume_jobs(extracted_text, topK)
    suggestions = build_resume_suggestions(profile, directions, matches)
    summary = build_resume_summary(extracted_text, profile, matches)

    return ResumeAnalysisResponse(
        fileName=file.filename or "resume",
        extractedText=compact(extracted_text, 6000),
        summary=summary,
        profile=profile,
        skills=skills,
        directions=directions,
        matches=matches,
        suggestions=suggestions,
        mode=extract_mode,
        configured=bool(api_key()),
    )


@app.post("/vision/analyze", response_model=VisionAnalyzeResponse)
async def analyze_vision(
    file: UploadFile = File(...),
    prompt: str = Form("请识别这张图片中的关键信息，并结合 Q_ITOffer 项目场景给出简洁说明。"),
    userContext: str = Form(""),
) -> VisionAnalyzeResponse:
    data = await checked_upload(file)
    if not api_key():
        return VisionAnalyzeResponse(
            fileName=file.filename or "image",
            answer="当前没有读取到 RAG_API_KEY，图片识别需要在 rag/.env 配置智谱 API key 后重启 RAG 服务。",
            mode="fallback",
            configured=False,
        )
    if not is_image_file(file.filename or "", file.content_type or ""):
        raise HTTPException(status_code=400, detail="图片识别只支持 PNG、JPG、JPEG、WEBP 等图片文件")
    try:
        vision_text = call_vision(vision_rag_prompt(prompt), file.content_type or "image/png", data)
        context = parse_user_context(userContext)
        user_docs = user_context_documents(context)
        image_doc = make_context_doc(
            "session:image",
            "上传图片识别内容",
            f"用户问题：{prompt}\n图片识别内容：{vision_text}",
            "personal_image",
        )
        extra_docs = [image_doc, *user_docs]
        match_text = "\n".join(doc.text for doc in extra_docs)
        matches = match_resume_jobs(match_text, 8)
        if matches:
            match_lines = ["根据上传图片和当前上下文匹配到的岗位："]
            for item in matches:
                match_lines.append(
                    f"- {item.title}；企业={item.company}；城市={item.city}；方向={item.category}；薪资={item.salary}；匹配度={item.score}%；原因={item.reason}"
                )
            extra_docs.append(make_context_doc("session:image-job-matches", "图片岗位匹配建议", "\n".join(match_lines), "personal_job_match"))
        hits = retrieve(f"{prompt}\n{vision_text}", 8, extra_docs)
        try:
            answer = call_vision_rag_llm(prompt, vision_text, hits, bool(user_docs))
        except requests.RequestException:
            answer = fallback_vision_rag_answer(prompt, vision_text, hits)
        return VisionAnalyzeResponse(
            fileName=file.filename or "image",
            answer=answer,
            mode="vision-rag",
            configured=True,
            sources=[source_from_hit(item) for item in hits],
        )
    except requests.RequestException as exc:
        return VisionAnalyzeResponse(
            fileName=file.filename or "image",
            answer=f"图片识别接口暂时不可用：{exc}",
            mode="fallback",
            configured=True,
        )


def rebuild_index(build_embeddings: bool) -> None:
    global INDEX
    docs = load_project_documents()
    docs.extend(load_database_documents())
    if not docs:
        docs.extend(load_seed_fallback_documents())
    for doc in docs:
        doc.terms = tokenize(doc.text + " " + doc.title)
        doc.text_hash = sha256(doc.title + "\n" + doc.text)
    hydrate_cached_embeddings(docs)
    if build_embeddings and api_key() and embed_model():
        refresh_embeddings(docs)
    INDEX = docs


def load_project_documents() -> list[Document]:
    docs: list[Document] = []
    roots = [
        PROJECT_ROOT / "README.md",
        PROJECT_ROOT / "docs",
        PROJECT_ROOT / "database",
        PROJECT_ROOT / "server" / "src",
        PROJECT_ROOT / "client" / "src",
        PROJECT_ROOT / "rag",
        PROJECT_ROOT / "scripts",
    ]
    for root in roots:
        paths = [root] if root.is_file() else sorted(root.rglob("*")) if root.exists() else []
        for path in paths:
            if not path.is_file() or should_skip_source(path):
                continue
            try:
                text = path.read_text(encoding="utf-8", errors="ignore")
            except OSError:
                continue
            text = redact_sensitive_text(text)
            rel = str(path.relative_to(PROJECT_ROOT))
            for index, chunk in enumerate(split_text(text, max_chars=1200), start=1):
                docs.append(Document(id=f"file:{rel}:{index}", title=f"{rel} #{index}", text=chunk, type="file", source=rel))
    return docs


def should_skip_source(path: Path) -> bool:
    rel_parts = set(path.relative_to(PROJECT_ROOT).parts)
    if rel_parts & SKIP_DIRS:
        return True
    if path.suffix.lower() in {".properties", ".local"}:
        return True
    if path.name.startswith(".env") or path.name.endswith(".lock"):
        return True
    if path.suffix.lower() not in TEXT_SUFFIXES:
        return True
    try:
        return path.stat().st_size > MAX_SOURCE_FILE_BYTES
    except OSError:
        return True


def redact_sensitive_text(text: str) -> str:
    replacements = [
        (r"(?i)(RAG_API_KEY\s*=\s*)[^\s#]+", r"\1[REDACTED]"),
        (r"(?i)(QITOFFER_DB_PASSWORD\s*=\s*)[^\s#]+", r"\1[REDACTED]"),
        (r"(?i)(VITE_AMAP_KEY\s*=\s*)[^\s#]+", r"\1[REDACTED]"),
        (r"(?i)(VITE_AMAP_SECURITY_JS_CODE\s*=\s*)[^\s#]+", r"\1[REDACTED]"),
        (r"(?i)(db\.password\s*=\s*)[^\s#]+", r"\1[REDACTED]"),
        (r"(?i)(--password=)[^\s\"',)]+", r"\1[REDACTED]"),
        (r"(?i)(IDENTIFIED BY\s+')[^']+(')", r"\1[REDACTED]\2"),
        (r"Qitoffer@2026", "[REDACTED]"),
    ]
    value = text
    for pattern, replacement in replacements:
        value = re.sub(pattern, replacement, value)
    return value


def load_database_documents() -> list[Document]:
    if not db_configured():
        return []
    try:
        conn = pymysql.connect(
            host=os.getenv("QITOFFER_DB_HOST", "127.0.0.1"),
            port=int(os.getenv("QITOFFER_DB_PORT", "3307")),
            user=os.getenv("QITOFFER_DB_USER", "qitoffer_app"),
            password=os.getenv("QITOFFER_DB_PASSWORD", "Qitoffer@2026"),
            database=os.getenv("QITOFFER_DB_NAME", "q_itoffer"),
            charset="utf8mb4",
            cursorclass=pymysql.cursors.DictCursor,
            connect_timeout=4,
            read_timeout=8,
        )
    except pymysql.MySQLError:
        return []

    docs: list[Document] = []
    with conn:
        for prefix, sql, formatter in [
            ("job", JOB_SQL, job_text),
            ("company", COMPANY_SQL, company_text),
        ]:
            try:
                docs.extend(query_docs(conn, prefix, sql, formatter))
            except pymysql.MySQLError:
                continue
    return docs


def query_docs(conn: pymysql.Connection, prefix: str, sql: str, formatter) -> list[Document]:
    docs: list[Document] = []
    with conn.cursor() as cursor:
        cursor.execute(sql)
        for row in cursor.fetchall():
            title, text, metadata = formatter(row)
            docs.append(Document(id=f"{prefix}:{row.get('id') or row.get('user_id')}", title=title, text=text, type=prefix, source="mysql:q_itoffer", metadata=metadata))
    return docs


def load_seed_fallback_documents() -> list[Document]:
    seed = PROJECT_ROOT / "database" / "seed.sql"
    if not seed.exists():
        return []
    text = seed.read_text(encoding="utf-8", errors="ignore")
    return [
        Document(id=f"seed:{index}", title=f"seed.sql #{index}", text=chunk, type="seed", source="database/seed.sql")
        for index, chunk in enumerate(split_text(text, max_chars=1200), start=1)
    ][:120]


def retrieve(question: str, top_k: int, extra_docs: list[Document] | None = None) -> list[tuple[Document, float]]:
    scored: dict[str, tuple[Document, float]] = {}
    candidate_docs = INDEX + (extra_docs or [])
    candidate_size = max(len(candidate_docs), top_k * 20, 80)
    for doc in extra_docs or []:
        scored[doc.id] = (doc, 0.9)
    for doc, score in lexical_retrieve(question, candidate_size, candidate_docs):
        scored[doc.id] = (doc, score)
    for doc, score in vector_retrieve(question, candidate_size):
        current = scored.get(doc.id)
        scored[doc.id] = (doc, max(score, current[1] if current else 0))
    scored = {doc_id: (doc, score + business_boost(question, doc)) for doc_id, (doc, score) in scored.items()}
    values = list(scored.values())
    values.sort(key=lambda item: item[1], reverse=True)
    return diversify_hits(values, top_k)


def diversify_hits(values: list[tuple[Document, float]], top_k: int) -> list[tuple[Document, float]]:
    selected: list[tuple[Document, float]] = []
    overflow: list[tuple[Document, float]] = []
    source_counts: dict[str, int] = {}
    max_per_source = 3 if top_k >= 6 else 2
    for item in values:
        source = item[0].source
        if source_counts.get(source, 0) < max_per_source:
            selected.append(item)
            source_counts[source] = source_counts.get(source, 0) + 1
        else:
            overflow.append(item)
        if len(selected) >= top_k:
            return selected
    for item in overflow:
        selected.append(item)
        if len(selected) >= top_k:
            break
    return selected


def lexical_retrieve(question: str, top_k: int, docs: list[Document] | None = None) -> list[tuple[Document, float]]:
    q_terms = tokenize(question)
    if not q_terms:
        return []
    scored: list[tuple[Document, float]] = []
    for doc in docs or INDEX:
        overlap = q_terms & doc.terms
        if not overlap:
            continue
        score = len(overlap) / max(len(q_terms), 1)
        if question in doc.text:
            score += 1.0
        if doc.type == "job":
            score += 0.08
        scored.append((doc, score))
    scored.sort(key=lambda item: item[1], reverse=True)
    return scored[:top_k]


def vector_retrieve(question: str, top_k: int) -> list[tuple[Document, float]]:
    docs = [doc for doc in INDEX if doc.embedding]
    if not docs or not api_key() or not embed_model():
        return []
    try:
        query_vector = call_embeddings([question])[0]
    except requests.RequestException:
        return []
    scored = [(doc, cosine(query_vector, doc.embedding or [])) for doc in docs]
    scored = [(doc, score) for doc, score in scored if score > 0]
    scored.sort(key=lambda item: item[1], reverse=True)
    return scored[:top_k]


def call_llm(question: str, hits: list[tuple[Document, float]], has_user_context: bool = False) -> str:
    context = "\n\n".join(f"[{index}] {doc.title}\n{compact(doc.text, 900)}" for index, (doc, _) in enumerate(hits, start=1))
    privacy_rule = (
        "本次请求包含当前登录用户的个人上下文，可以使用来源为 session:current-user 的简历和投递记录。"
        if has_user_context
        else "本次请求没有当前登录用户的个人上下文。禁止使用、猜测或编造任何个人简历、姓名、投递记录、申请进度；如果用户问自己的简历或投递情况，必须提示先登录或上传简历。"
    )
    response = requests.post(
        f"{base_url()}/chat/completions",
        headers=auth_headers(),
        json={
            "model": chat_model(),
            "temperature": 0.2,
            "messages": [
                {
                    "role": "system",
                    "content": f"你是 Q_ITOffer 项目的内部问答助手。必须回答用户提出的原问题，不要改写成另一个问题；只根据给定上下文回答，答案要简洁，并说明可引用的来源序号。{privacy_rule}",
                },
                {"role": "user", "content": f"问题：{question}\n\n上下文：\n{context}"},
            ],
        },
        timeout=30,
    )
    response.raise_for_status()
    payload = response.json()
    return payload["choices"][0]["message"]["content"].strip()


def fallback_answer(question: str, hits: list[tuple[Document, float]]) -> str:
    lines = [f"关于“{question}”，当前知识库里最相关的信息如下："]
    for index, (doc, _) in enumerate(hits[:4], start=1):
        lines.append(f"{index}. {doc.title}：{compact(doc.text, 180)}")
    lines.append("这是本地检索模式的回答。配置 RAG_API_KEY 并重启服务后，会自动切换为大模型整合回答。")
    return "\n".join(lines)


def source_from_hit(hit: tuple[Document, float]) -> Source:
    doc, score = hit
    return Source(id=doc.id, title=doc.title, type=doc.type, source=doc.source, snippet=compact(doc.text, 150), score=round(score, 4))


def business_boost(question: str, doc: Document) -> float:
    query = question.lower()
    source = doc.source.replace("/", "\\").lower()
    boost = 0.0
    if doc.type.startswith("personal_"):
        boost += 1.0
    asks_admin = any(word in question for word in ["后台", "管理端", "管理系统", "管理员"]) or "admin" in query
    asks_admin_api = asks_admin and ("\u63a5\u53e3" in question or "api" in query or "servlet" in query)
    asks_resume = "简历" in question or "resume" in query
    asks_resume_document = asks_resume and any(word in question for word in ["\u9644\u4ef6", "\u4e0a\u4f20", "\u6587\u6863"]) or "document" in query or "upload" in query
    asks_application = any(word in question for word in ["投递", "申请"]) or "application" in query
    asks_map = "地图" in question or "热力" in question or "amap" in query

    if asks_admin:
        if asks_admin_api and "adminapiservlet" in source:
            boost += 1.0
        if "client\\src\\admin" in source:
            boost += 0.35 if asks_admin_api else 0.5
        if "adminapiservlet" in source:
            boost += 0.48
        if "web-inf\\jsp\\manage" in source or "servlet\\manage" in source:
            boost += 0.08
        if doc.type == "job":
            boost -= 0.08
    if asks_resume:
        if asks_resume_document and "resumeapiservlet" in source:
            boost += 0.75
        if asks_resume_document and "jobapiservlet" in source:
            boost += 0.55
        if asks_resume_document and ("client\\src\\pages\\resume" in source or "client\\src\\pages\\jobdetail" in source):
            boost += 0.42
        if "web-inf\\jsp\\manage" in source:
            boost -= 0.18
        if "resume" in source or doc.type == "resume":
            boost += 0.35
    if asks_application:
        if "application" in source or "jobapiservlet" in source or doc.type == "application":
            boost += 0.32
    if asks_map:
        if "mapheat" in source or "statsapiservlet" in source or "amap" in source:
            boost += 0.35
    return boost


def needs_private_context(question: str) -> bool:
    private_patterns = [
        "我的简历",
        "我的投递",
        "我的申请",
        "我的岗位",
        "我的技能",
        "我的经历",
        "申请进度",
        "投递记录",
        "我适合",
        "适合我",
        "适合什么岗位",
        "适合哪些岗位",
        "我能干",
        "能干什么岗位",
        "我能做",
        "能做什么岗位",
        "我可以做",
        "给我推荐",
        "推荐给我",
        "根据我的",
        "我的信息",
        "我的个人",
        "我投递",
        "我已经投",
    ]
    return any(pattern in question for pattern in private_patterns)


def is_personal_job_question(question: str) -> bool:
    patterns = [
        "我能干",
        "我能做",
        "我可以做",
        "适合什么岗位",
        "适合哪些岗位",
        "能干什么岗位",
        "能做什么岗位",
        "给我推荐岗位",
        "推荐给我岗位",
    ]
    return any(pattern in question for pattern in patterns)


def personal_job_answer(user_docs: list[Document]) -> str:
    resume_text = "\n".join(doc.text for doc in user_docs if doc.type == "personal_resume")
    match_doc = next((doc for doc in user_docs if doc.type == "personal_job_match"), None)
    skills = extract_skills(resume_text)
    directions = infer_directions(resume_text, skills)
    lines = ["根据你当前登录账号的简历内容，我建议优先看这些方向："]
    if directions:
        lines.append("、".join(f"{index}. {item}" for index, item in enumerate(directions, start=1)))
    if skills:
        lines.append(f"我检测到的关键词包括：{'、'.join(skills[:12])}。")
    if match_doc:
        recommendations = [
            line[2:].strip()
            for line in match_doc.text.splitlines()
            if line.startswith("- ")
        ][:5]
        if recommendations:
            lines.append("系统岗位库里较匹配的岗位有：")
            lines.extend(f"{index}. {item}" for index, item in enumerate(recommendations, start=1))
    if not skills:
        lines.append("当前简历里的技能关键词还不够清晰，可以先补充技术栈、项目经历和目标岗位，再让我重新匹配。")
    lines.append("这次只使用当前登录人的简历草稿/已保存简历和公共岗位库，没有读取其他用户的个人信息。")
    return "\n".join(lines)


def user_context_documents(context: dict[str, Any] | None) -> list[Document]:
    if not isinstance(context, dict):
        return []
    user = context.get("user") if isinstance(context.get("user"), dict) else {}
    user_name = str(user.get("fullName") or user.get("username") or "当前用户").strip()
    docs: list[Document] = []

    resume = context.get("resume") if isinstance(context.get("resume"), dict) else {}
    if resume:
        lines = [
            f"当前登录用户：{user_name}",
            f"账号角色：{user.get('role') or ''}",
            f"邮箱：{user.get('email') or resume.get('email') or ''}",
            f"电话：{user.get('phone') or resume.get('phone') or ''}",
            f"学历：{resume.get('education') or ''}",
            f"专业：{resume.get('major') or ''}",
            f"期望城市：{resume.get('expected_city') or resume.get('expectedCity') or ''}",
            f"期望薪资：{resume.get('expected_salary') or resume.get('expectedSalary') or ''}",
            f"技能：{resume.get('skills') or ''}",
            f"自我介绍：{resume.get('self_intro') or resume.get('selfIntro') or ''}",
        ]
        modules = context.get("resumeModules") if isinstance(context.get("resumeModules"), dict) else {}
        lines.extend(format_resume_modules(modules))
        resume_text = "\n".join(lines)
        docs.append(make_context_doc("session:resume", "当前登录用户简历", resume_text, "personal_resume"))
        matches = match_resume_jobs(resume_text, 8)
        if matches:
            match_lines = [f"当前登录用户：{user_name}", "根据当前登录用户简历匹配到的岗位："]
            for item in matches:
                match_lines.append(
                    f"- {item.title}；企业={item.company}；城市={item.city}；方向={item.category}；薪资={item.salary}；匹配度={item.score}%；原因={item.reason}"
                )
            docs.append(make_context_doc("session:job-matches", "当前登录用户岗位匹配建议", "\n".join(match_lines), "personal_job_match"))

    documents = context.get("resumeDocuments")
    if isinstance(documents, list):
        for index, item in enumerate(documents[:8], start=1):
            if not isinstance(item, dict):
                continue
            text = "\n".join(
                [
                    f"当前登录用户：{user_name}",
                    f"简历附件：{item.get('originalFilename') or item.get('original_filename') or ''}",
                    f"解析文本：{compact(str(item.get('parsedText') or item.get('parsed_text') or ''), 1600)}",
                    f"分析结果：{compact(str(item.get('analysisJson') or item.get('analysis_json') or ''), 1200)}",
                ]
            )
            docs.append(make_context_doc(f"session:resume-document:{index}", "当前登录用户简历附件", text, "personal_resume"))

    applications = context.get("applications")
    if isinstance(applications, list) and applications:
        lines = [f"当前登录用户：{user_name}", "当前登录用户的职位申请记录："]
        for item in applications[:30]:
            if not isinstance(item, dict):
                continue
            lines.append(
                "；".join(
                    [
                        f"岗位={item.get('title') or ''}",
                        f"企业={item.get('companyName') or item.get('company_name') or ''}",
                        f"城市={item.get('city') or ''}",
                        f"状态={item.get('status') or ''}",
                        f"面试反馈={item.get('interviewResponse') or item.get('interview_response') or ''}",
                        f"投递时间={item.get('appliedAt') or item.get('applied_at') or ''}",
                        f"简历附件={item.get('resumeFilename') or item.get('resume_filename') or ''}",
                    ]
                )
            )
        docs.append(make_context_doc("session:applications", "当前登录用户投递记录", "\n".join(lines), "personal_application"))

    return docs


def format_resume_modules(modules: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    labels = {
        "educations": "教育经历",
        "experiences": "工作/实习经历",
        "projects": "项目经历",
        "skills": "技能标签",
        "certificates": "证书/培训",
    }
    for key, label in labels.items():
        items = modules.get(key)
        if not isinstance(items, list) or not items:
            continue
        lines.append(f"{label}：")
        for item in items[:12]:
            if not isinstance(item, dict):
                continue
            parts = []
            for value_key in [
                "school",
                "degree",
                "major",
                "company",
                "position",
                "name",
                "role_name",
                "roleName",
                "tech_stack",
                "techStack",
                "level_name",
                "levelName",
                "issuer",
                "description",
            ]:
                value = item.get(value_key)
                if value:
                    parts.append(str(value))
            if parts:
                lines.append(f"- {'；'.join(parts)}")
    return lines


def make_context_doc(doc_id: str, title: str, text: str, doc_type: str) -> Document:
    doc = Document(id=doc_id, title=title, text=text, type=doc_type, source="session:current-user")
    doc.terms = tokenize(doc.text + " " + doc.title)
    doc.text_hash = sha256(doc.title + "\n" + doc.text)
    return doc


async def checked_upload(file: UploadFile) -> bytes:
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="上传文件为空")
    max_mb = int(os.getenv("RAG_MAX_UPLOAD_MB", "10"))
    if len(data) > max_mb * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"上传文件不能超过 {max_mb}MB")
    return data


def extract_resume_text(filename: str, content_type: str, data: bytes) -> tuple[str, str]:
    lower = filename.lower()
    if lower.endswith(".txt") or content_type.startswith("text/"):
        return data.decode("utf-8", errors="ignore"), "text"
    if lower.endswith(".pdf") or content_type == "application/pdf":
        return extract_pdf_text(data), "pdf"
    if lower.endswith(".docx") or "wordprocessingml.document" in content_type:
        return extract_docx_text(data), "docx"
    if is_image_file(filename, content_type):
        return "", "vision"
    raise HTTPException(status_code=400, detail="简历上传仅支持 PDF、DOCX、TXT、PNG、JPG、JPEG、WEBP")


def extract_pdf_text(data: bytes) -> str:
    try:
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 - return a user-readable failure to the RAG panel.
        return f"PDF 文本解析失败：{exc}"


def extract_docx_text(data: bytes) -> str:
    try:
        from io import BytesIO

        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(data))
        paragraphs = [item.text for item in doc.paragraphs if item.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(paragraphs)
    except Exception as exc:  # noqa: BLE001
        return f"DOCX 文本解析失败：{exc}"


def vision_extract_resume(filename: str, content_type: str, data: bytes) -> str:
    if not api_key():
        return ""
    prompt = (
        "这是一份图片格式的个人简历。请提取姓名、邮箱、手机号、学历、专业、技能、项目经历、工作/实习经历、"
        "求职方向等信息，输出为紧凑中文文本，方便招聘系统做岗位匹配。"
    )
    try:
        return call_vision(prompt, content_type or mimetypes.guess_type(filename)[0] or "image/png", data)
    except requests.RequestException:
        return ""


def call_vision(prompt: str, content_type: str, data: bytes) -> str:
    image_url = f"data:{content_type};base64,{base64.b64encode(data).decode('ascii')}"
    response = requests.post(
        f"{base_url()}/chat/completions",
        headers=auth_headers(),
        json={
            "model": vision_model(),
            "temperature": 0.2,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": image_url}},
                    ],
                }
            ],
        },
        timeout=45,
    )
    response.raise_for_status()
    payload = response.json()
    return payload["choices"][0]["message"]["content"].strip()


def vision_rag_prompt(user_prompt: str) -> str:
    return (
        "请完整识别图片里的求职、简历、项目、技能、证书、岗位或截图内容。"
        "如果图片像简历，请提取姓名、学历、专业、技能、项目经历、工作经历、证书和求职方向；"
        "如果图片像岗位或项目材料，请提取技术栈、职责、成果和关键词。"
        f"用户接下来要你结合 Q_ITOffer 岗位库回答的问题是：{user_prompt}"
    )


def parse_user_context(raw: str) -> dict[str, Any] | None:
    if not raw.strip():
        return None
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return None
    return data if isinstance(data, dict) else None


def call_vision_rag_llm(prompt: str, vision_text: str, hits: list[tuple[Document, float]], has_user_context: bool) -> str:
    context = "\n\n".join(f"[{index}] {doc.title}\n{compact(doc.text, 1000)}" for index, (doc, _) in enumerate(hits, start=1))
    privacy_rule = (
        "本次请求包含当前登录用户上下文，可以结合 session:current-user 的简历草稿、已保存简历和投递记录。"
        if has_user_context
        else "本次请求没有当前登录用户上下文，只能使用图片识别内容和公共岗位库，不能编造用户个人简历或投递记录。"
    )
    response = requests.post(
        f"{base_url()}/chat/completions",
        headers=auth_headers(),
        json={
            "model": chat_model(),
            "temperature": 0.25,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "你是 Q_ITOffer 的多模态 RAG 求职助手。必须回答用户原问题，不能只复述图片内容。"
                        "请结合图片识别结果、当前登录用户上下文和岗位库来源，给出类似 LLM 助手的结构化建议："
                        "先说结论，再列适合岗位方向、依据、可投递岗位示例和提升建议，并引用来源序号。"
                        f"{privacy_rule}"
                    ),
                },
                {
                    "role": "user",
                    "content": f"用户问题：{prompt}\n\n图片识别结果：\n{vision_text}\n\nRAG 上下文：\n{context}",
                },
            ],
        },
        timeout=45,
    )
    response.raise_for_status()
    payload = response.json()
    return payload["choices"][0]["message"]["content"].strip()


def fallback_vision_rag_answer(prompt: str, vision_text: str, hits: list[tuple[Document, float]]) -> str:
    lines = [
        f"我已识别图片，并结合 RAG 上下文回答你的问题：{prompt}",
        "",
        "图片关键信息：",
        compact(vision_text, 700),
    ]
    match_docs = [doc for doc, _ in hits if doc.type == "personal_job_match"]
    if match_docs:
        lines.append("")
        lines.append("匹配到的岗位建议：")
        for doc in match_docs[:2]:
            for line in doc.text.splitlines():
                if line.startswith("- "):
                    lines.append(line)
    lines.append("")
    lines.append("这次回答使用了图片识别内容、当前登录用户上下文和公共岗位库；未登录时不会读取任何其他用户个人信息。")
    return "\n".join(lines)


def extract_resume_profile(text: str) -> dict[str, Any]:
    clean = re.sub(r"\s+", " ", text).strip()
    email = first_match(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", clean)
    phone = first_match(r"(?:\+?86[- ]?)?1[3-9]\d{9}", clean)
    name = infer_name(text)
    education = first_keyword(clean, ["博士", "硕士", "研究生", "本科", "大专", "专科"])
    skills = extract_skills(clean)
    return {
        "fullName": name,
        "email": email,
        "phone": phone,
        "education": education,
        "major": infer_major(clean),
        "expectedCity": infer_city(clean),
        "skills": skills,
        "selfIntro": compact(clean, 280),
    }


def extract_skills(text: str) -> list[str]:
    known = [
        "Java",
        "Spring",
        "Spring Boot",
        "Servlet",
        "JSP",
        "Go",
        "Golang",
        "MySQL",
        "Redis",
        "Python",
        "FastAPI",
        "RAG",
        "LangChain",
        "Agent",
        "AIGC",
        "LLM",
        "多模态",
        "React",
        "Vue",
        "TypeScript",
        "JavaScript",
        "HTML",
        "CSS",
        "Node",
        "Docker",
        "Linux",
        "Nginx",
        "Git",
        "数据分析",
        "机器学习",
        "深度学习",
        "测试",
        "自动化测试",
        "UI",
        "产品",
    ]
    lower = text.lower()
    found = []
    for skill in known:
        if skill.lower() in lower or skill in text:
            found.append(skill)
    return found[:18]


def infer_directions(text: str, skills: list[str]) -> list[str]:
    lower = (text + " " + " ".join(skills)).lower()
    rules = [
        ("前端开发", ["react", "vue", "typescript", "javascript", "html", "css"]),
        ("Java开发", ["java", "spring", "servlet", "jsp"]),
        ("Go后端开发", ["go", "golang"]),
        ("Python开发", ["python", "fastapi", "django", "flask"]),
        ("RAG/大模型应用开发", ["rag", "langchain", "agent", "aigc", "llm", "多模态", "大模型"]),
        ("数据分析", ["数据分析", "sql", "python", "机器学习"]),
        ("测试工程师", ["测试", "自动化测试", "selenium", "pytest"]),
        ("产品经理", ["产品", "原型", "需求"]),
        ("UI设计", ["ui", "figma", "视觉"]),
    ]
    directions = [name for name, words in rules if any(word.lower() in lower for word in words)]
    if not directions:
        directions = ["IT 综合岗位"]
    return directions[:5]


def match_resume_jobs(text: str, top_k: int) -> list[JobMatch]:
    if not text.strip():
        return []
    q_terms = tokenize(text)
    hits: list[tuple[Document, float]] = []
    for doc in INDEX:
        if doc.type != "job":
            continue
        overlap = q_terms & doc.terms
        score = len(overlap) / max(len(q_terms), 1)
        lower_text = text.lower()
        for keyword in ["java", "spring", "react", "typescript", "python", "mysql", "go", "golang", "rag", "langchain", "agent", "aigc", "大模型", "多模态", "测试", "产品", "ui"]:
            if keyword in lower_text and keyword in doc.text.lower():
                score += 0.06
        if score > 0:
            hits.append((doc, score))
    hits.sort(key=lambda item: item[1], reverse=True)
    matches: list[JobMatch] = []
    seen: set[int] = set()
    for doc, score in hits:
        job_id = int(doc.metadata.get("jobId") or 0)
        if job_id in seen:
            continue
        seen.add(job_id)
        matches.append(
            JobMatch(
                jobId=job_id,
                title=str(doc.metadata.get("title") or doc.title),
                company=str(doc.metadata.get("company") or ""),
                city=str(doc.metadata.get("city") or ""),
                category=str(doc.metadata.get("category") or ""),
                salary=str(doc.metadata.get("salary") or ""),
                score=round(min(score, 1.0) * 100, 1),
                reason=compact(doc.text, 120),
            )
        )
        if len(matches) >= top_k:
            break
    return matches


def build_resume_suggestions(profile: dict[str, Any], directions: list[str], matches: list[JobMatch]) -> list[str]:
    suggestions: list[str] = []
    if not profile.get("email") or not profile.get("phone"):
        suggestions.append("补充邮箱和手机号，方便企业联系。")
    if not profile.get("skills"):
        suggestions.append("补充技能关键词，例如 Java、React、MySQL、Python 等。")
    if not directions:
        suggestions.append("在简历中明确目标岗位方向，匹配效果会更稳定。")
    if not matches:
        suggestions.append("当前岗位匹配较少，可以补充项目经历、技术栈和期望城市。")
    suggestions.append("投递前建议把项目经历写成“背景-职责-技术-结果”的结构。")
    return suggestions[:5]


def build_resume_summary(text: str, profile: dict[str, Any], matches: list[JobMatch]) -> str:
    if not text.strip():
        return "暂未从文件中提取到足够文本。如果是图片简历，请确认已配置智谱视觉模型。"
    skill_text = "、".join(profile.get("skills") or []) or "技能待补充"
    match_text = f"推荐优先查看 {matches[0].title} 等岗位。" if matches else "暂未匹配到明显岗位。"
    return f"简历已解析，识别到技能：{skill_text}。{match_text}"


def hydrate_cached_embeddings(docs: list[Document]) -> None:
    try:
        conn = embedding_conn()
    except sqlite3.Error:
        return
    with conn:
        ensure_embedding_schema(conn)
        for doc in docs:
            row = conn.execute("SELECT vector_json FROM embeddings WHERE text_hash = ?", (doc.text_hash,)).fetchone()
            if row:
                try:
                    doc.embedding = json.loads(row[0])
                except json.JSONDecodeError:
                    doc.embedding = None


def refresh_embeddings(docs: list[Document]) -> None:
    global LAST_EMBEDDING_ERROR
    LAST_EMBEDDING_ERROR = ""
    pending = [doc for doc in docs if not doc.embedding and len(doc.text) > 20]
    if not pending:
        return
    try:
        conn = embedding_conn()
        ensure_embedding_schema(conn)
        for index in range(0, len(pending), 24):
            batch = pending[index:index + 24]
            vectors = call_embeddings([compact(doc.text, 1800) for doc in batch])
            with conn:
                for doc, vector in zip(batch, vectors):
                    doc.embedding = vector
                    conn.execute(
                        "INSERT OR REPLACE INTO embeddings (text_hash, vector_json, updated_at) VALUES (?, ?, CURRENT_TIMESTAMP)",
                        (doc.text_hash, json.dumps(vector)),
                    )
    except (requests.RequestException, sqlite3.Error, KeyError, IndexError, TypeError) as exc:
        LAST_EMBEDDING_ERROR = str(exc)


def call_embeddings(texts: list[str]) -> list[list[float]]:
    response = requests.post(
        f"{base_url()}/embeddings",
        headers=auth_headers(),
        json={"model": embed_model(), "input": texts},
        timeout=45,
    )
    try:
        response.raise_for_status()
    except requests.HTTPError as exc:
        raise requests.HTTPError(f"{exc}; response={response.text[:500]}", response=response) from exc
    payload = response.json()
    data = sorted(payload["data"], key=lambda item: item.get("index", 0))
    return [item["embedding"] for item in data]


def embedding_conn() -> sqlite3.Connection:
    path = Path(os.getenv("RAG_INDEX_DB", RAG_ROOT / "index.sqlite3"))
    path.parent.mkdir(parents=True, exist_ok=True)
    return sqlite3.connect(path)


def ensure_embedding_schema(conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS embeddings (
          text_hash TEXT PRIMARY KEY,
          vector_json TEXT NOT NULL,
          updated_at TEXT NOT NULL
        )
        """
    )


def cosine(left: list[float], right: list[float]) -> float:
    if not left or not right or len(left) != len(right):
        return 0
    dot = sum(a * b for a, b in zip(left, right))
    left_norm = math.sqrt(sum(a * a for a in left))
    right_norm = math.sqrt(sum(b * b for b in right))
    if left_norm == 0 or right_norm == 0:
        return 0
    return dot / (left_norm * right_norm)


def tokenize(text: str) -> set[str]:
    lower = text.lower()
    ascii_terms = re.findall(r"[a-z0-9_+#.]{2,}", lower)
    for keyword, aliases in TOKEN_ALIASES.items():
        if keyword in text:
            ascii_terms.extend(aliases)
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", text)
    chinese_bigrams = ["".join(chinese_chars[index:index + 2]) for index in range(len(chinese_chars) - 1)]
    chinese_trigrams = ["".join(chinese_chars[index:index + 3]) for index in range(len(chinese_chars) - 2)]
    return set(ascii_terms + chinese_bigrams + chinese_trigrams)


def split_text(text: str, max_chars: int = 900) -> list[str]:
    blocks = [block.strip() for block in re.split(r"\n{2,}", text) if block.strip()]
    chunks: list[str] = []
    current = ""
    for block in blocks:
        if len(block) > max_chars:
            if current:
                chunks.append(current)
                current = ""
            chunks.extend(block[index:index + max_chars] for index in range(0, len(block), max_chars))
            continue
        if len(current) + len(block) + 2 > max_chars and current:
            chunks.append(current)
            current = block
        else:
            current = f"{current}\n\n{block}".strip()
    if current:
        chunks.append(current)
    return chunks


def compact(text: str, limit: int) -> str:
    value = re.sub(r"\s+", " ", text).strip()
    return value if len(value) <= limit else value[:limit].rstrip() + "..."


def first_match(pattern: str, text: str) -> str:
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def first_keyword(text: str, keywords: list[str]) -> str:
    return next((word for word in keywords if word in text), "")


def infer_name(text: str) -> str:
    for line in text.splitlines()[:8]:
        clean = re.sub(r"\s+", "", line)
        if 2 <= len(clean) <= 6 and re.fullmatch(r"[\u4e00-\u9fffA-Za-z]+", clean):
            return clean
        match = re.search(r"(?:姓名|Name)[:： ]*([\u4e00-\u9fffA-Za-z]{2,20})", line, re.IGNORECASE)
        if match:
            return match.group(1)
    return ""


def infer_major(text: str) -> str:
    match = re.search(r"(?:专业|Major)[:： ]*([\u4e00-\u9fffA-Za-z0-9+\- ]{2,40})", text, re.IGNORECASE)
    return match.group(1).strip() if match else ""


def infer_city(text: str) -> str:
    cities = ["北京", "上海", "广州", "深圳", "杭州", "南京", "成都", "武汉", "西安", "长沙", "郑州", "重庆", "苏州"]
    return next((city for city in cities if city in text), "")


def is_image_file(filename: str, content_type: str) -> bool:
    lower = filename.lower()
    return content_type.startswith("image/") or lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp"))


def sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def api_key() -> str:
    return os.getenv("RAG_API_KEY", "").strip()


def base_url() -> str:
    return os.getenv("RAG_BASE_URL", "https://open.bigmodel.cn/api/paas/v4").rstrip("/")


def chat_model() -> str:
    return os.getenv("RAG_CHAT_MODEL", DEFAULT_CHAT_MODEL).strip()


def embed_model() -> str:
    return os.getenv("RAG_EMBED_MODEL", DEFAULT_EMBED_MODEL).strip()


def vision_model() -> str:
    return os.getenv("RAG_VISION_MODEL", DEFAULT_VISION_MODEL).strip()


def auth_headers() -> dict[str, str]:
    return {"Authorization": f"Bearer {api_key()}", "Content-Type": "application/json"}


def allowed_origins() -> list[str]:
    raw = os.getenv("RAG_ALLOWED_ORIGINS", "*").strip()
    return ["*"] if not raw else [item.strip() for item in raw.split(",") if item.strip()]


def db_configured() -> bool:
    return bool(os.getenv("QITOFFER_DB_HOST", "127.0.0.1"))


def job_text(row: dict[str, Any]) -> tuple[str, str, dict[str, Any]]:
    title = f"岗位：{row.get('title')} / {row.get('city')}"
    salary = f"{row.get('salary_min')}-{row.get('salary_max')}"
    text = (
        f"岗位 {row.get('title')}，城市 {row.get('city')}，方向 {row.get('category')}，"
        f"薪资 {salary}，企业 {row.get('company_name')}，学历 {row.get('education')}，经验 {row.get('experience')}，"
        f"亮点 {row.get('highlights')}。描述：{row.get('description')} 要求：{row.get('requirement_text')}"
    )
    metadata = {
        "jobId": row.get("id"),
        "title": row.get("title"),
        "city": row.get("city"),
        "category": row.get("category"),
        "company": row.get("company_name"),
        "salary": salary,
    }
    return title, text, metadata


def company_text(row: dict[str, Any]) -> tuple[str, str, dict[str, Any]]:
    title = f"企业：{row.get('name')}"
    text = (
        f"企业 {row.get('name')}，城市 {row.get('city')}，行业 {row.get('industry')}，"
        f"规模 {row.get('scale')}，融资 {row.get('financing_stage')}，评分 {row.get('rating')}。介绍：{row.get('description')}"
    )
    return title, text, {"companyId": row.get("id"), "city": row.get("city"), "industry": row.get("industry")}


def resume_text(row: dict[str, Any]) -> tuple[str, str, dict[str, Any]]:
    title = f"简历：{row.get('full_name')}"
    text = (
        f"用户 {row.get('username')}，姓名 {row.get('full_name')}，学历 {row.get('education')}，"
        f"专业 {row.get('major')}，期望城市 {row.get('expected_city')}，期望薪资 {row.get('expected_salary')}，"
        f"技能 {row.get('skills')}。自我介绍：{row.get('self_intro')}。附件简历：{row.get('document_text') or ''}"
    )
    return title, text, {"userId": row.get("user_id")}


def application_text(row: dict[str, Any]) -> tuple[str, str, dict[str, Any]]:
    title = f"投递：{row.get('full_name')} -> {row.get('title')}"
    text = (
        f"{row.get('full_name')} 投递 {row.get('company_name')} 的 {row.get('title')}，"
        f"状态 {row.get('status')}，面试响应 {row.get('interview_response')}，留言 {row.get('message')}，"
        f"简历附件 {row.get('resume_filename') or '无'}，投递时间 {row.get('applied_at')}。"
    )
    return title, text, {"applicationId": row.get("id")}


JOB_SQL = """
SELECT j.id, j.title, j.category, j.salary_min, j.salary_max, j.city, j.education, j.experience,
       j.highlights, j.description, j.requirement_text, c.name AS company_name
FROM jobs j
JOIN companies c ON c.id = j.company_id
WHERE j.status = 'OPEN'
ORDER BY j.posted_at DESC
LIMIT 1000
"""

COMPANY_SQL = """
SELECT id, name, city, industry, scale, financing_stage, rating, description
FROM companies
ORDER BY rating DESC, id DESC
LIMIT 400
"""

RESUME_SQL = """
SELECT u.id AS user_id, u.username, u.full_name, p.education, p.major, p.expected_city,
       p.expected_salary, p.skills, p.self_intro,
       GROUP_CONCAT(rd.parsed_text SEPARATOR '\n') AS document_text
FROM users u
LEFT JOIN applicant_profiles p ON p.user_id = u.id
LEFT JOIN resume_documents rd ON rd.user_id = u.id
WHERE u.role = 'APPLICANT'
GROUP BY u.id, u.username, u.full_name, p.education, p.major, p.expected_city, p.expected_salary, p.skills, p.self_intro
ORDER BY u.id DESC
LIMIT 400
"""

APPLICATION_SQL = """
SELECT a.id, a.status, a.interview_response, a.message, a.applied_at,
       rd.original_filename AS resume_filename,
       u.full_name, j.title, c.name AS company_name
FROM applications a
JOIN users u ON u.id = a.applicant_id
JOIN jobs j ON j.id = a.job_id
JOIN companies c ON c.id = j.company_id
LEFT JOIN resume_documents rd ON rd.id = a.resume_document_id
ORDER BY a.applied_at DESC
LIMIT 700
"""
